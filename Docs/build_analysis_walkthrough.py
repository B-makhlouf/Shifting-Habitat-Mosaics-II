from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = r"Docs\SMH2_analysis_walkthrough.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 90, 90)
PALE = "E8EEF5"


def font(run, size=11, bold=False, italic=False, color=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color


def shade(paragraph, fill):
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


doc = Document()
sec = doc.sections[0]
sec.top_margin = Inches(0.75)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(0.85)
sec.right_margin = Inches(0.85)
sec.header_distance = Inches(0.492)
sec.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in [
    ("Title", 24, DARK, 0, 5),
    ("Heading 1", 16, BLUE, 14, 6),
    ("Heading 2", 12.5, BLUE, 9, 4),
]:
    s = styles[name]
    s.font.name = "Calibri"
    s._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    s._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    s.font.size = Pt(size)
    s.font.color.rgb = color
    s.font.bold = True
    s.paragraph_format.space_before = Pt(before)
    s.paragraph_format.space_after = Pt(after)
    s.paragraph_format.keep_with_next = True

lb = styles["List Bullet"]
lb.font.name = "Calibri"
lb.font.size = Pt(10.5)
lb.paragraph_format.left_indent = Inches(0.38)
lb.paragraph_format.first_line_indent = Inches(-0.19)
lb.paragraph_format.space_after = Pt(4)
lb.paragraph_format.line_spacing = 1.12

header = sec.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
font(header.add_run("Shifting Habitat Mosaics II  |  analysis drafting guide"),
     size=8.5, color=MUTED)

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
font(footer.add_run("Internal drafting aid - verify parameter labels before submission"),
     size=8, color=MUTED, italic=True)

p = doc.add_paragraph(style="Title")
p.add_run("Analysis Walkthrough for Methods and Results")
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
font(p.add_run("Yukon and Kuskokwim Chinook salmon natal-origin assignments, "
               "landscape contours, and variance buffering"), size=12, color=MUTED)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(10)
shade(p, PALE)
r = p.add_run("Scope. ")
font(r, bold=True, color=DARK)
font(p.add_run("This is a concise reconstruction of the current scripts and saved outputs. "
               "It is intended as a checklist for manuscript drafting, not final Methods prose."),
     size=10.5)

doc.add_heading("1. Initial natal-origin assignments", level=1)

items = [
    ("Inputs and years",
     "Combined otolith natal-isotope observations with daily CPUE weighting and basin-scale "
     "isoscapes for Kuskokwim (2017-2022) and Yukon (2015, 2016, 2018, 2021)."),
    ("Sampling correction",
     "Divided each run into five day-of-year strata and calculated a stratum weight as the "
     "summed daily CPUE proportion divided by the summed otolith sampling proportion. "
     "The current assignment matrix is then summed across fish; confirm whether the stratum "
     "weight is intended to enter the final assignment calculation, because it is calculated "
     "in the script but is not visibly multiplied into A."),
    ("Likelihood",
     "For every fish and eligible river reach, evaluated a Gaussian likelihood comparing the "
     "fish's natal 87Sr/86Sr value with the reach-level predicted isoscape value. Total error "
     "combined isoscape uncertainty, otolith measurement uncertainty, and water-sample uncertainty "
     "in quadrature."),
    ("Spatial priors",
     "Excluded reaches below Strahler order 3. Kuskokwim assignments also used the UniPh2oNoE "
     "habitat weight and removed selected high-order reaches without spawning presence. Yukon "
     "assignments excluded channel slope >2, used spawning-presence restrictions in the Lower and "
     "Middle regions, and multiplied assignments by genetic probabilities for Lower, Middle, and "
     "Upper Yukon; incomplete genetics were replaced by equal one-third regional weights."),
    ("Normalization and thresholding",
     "Normalized each fish's reach probabilities to sum to one, rescaled them relative to that "
     "fish's maximum assignment, and set values below the selected sensitivity threshold to zero. "
     "Annual reach scores were summed, converted to basin shares, and multiplied by total run size "
     "to estimate numbers of fish."),
    ("Map outputs",
     "Mapped each reach's proportional share of annual production, with line width controlled by "
     "stream order and color showing production relative to the annual maximum."),
]
for label, text in items:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(label + ": ")
    font(r, bold=True, color=DARK)
    font(p.add_run(text))

doc.add_heading("2. Contour figures: the landscape distribution of production", level=1)
items = [
    ("Contour input",
     "Re-ran the assignment calculation for the contour pass at tau = 0.9 and read those "
     "thresholded reach files. The separate contour filter is currently assignment_norm > 0.0."),
    ("Landscape variables",
     "Joined reach assignments to watershed slope and distance upstream. Plotted log10 watershed "
     "slope on the x-axis and distance upstream on the y-axis (displayed in 100-km units)."),
    ("Weighted two-dimensional density",
     "Estimated a bivariate kernel density with ks::kde, using assignment_norm as the reach weight "
     "and a plug-in bandwidth matrix. Filled contours represent weighted density quantiles "
     "(10, 20, 40, 60, 80, and 90%)."),
    ("Cross-year reference",
     "Pooled annual data within each basin after giving every year equal total weight, then drew "
     "the pooled 80% highest-density boundary and weighted center as a fixed reference for judging "
     "annual shifts."),
    ("Interpretation",
     "Treat contour lobes as recurring combinations of geomorphic setting and upstream position "
     "associated with assigned production. They summarize a weighted landscape distribution; "
     "they are not confidence regions for individual fish."),
]
for label, text in items:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(label + ": ")
    font(r, bold=True, color=DARK)
    font(p.add_run(text))

doc.add_heading("3. Variance analysis and portfolio-effect test", level=1)
items = [
    ("Annual production series",
     "Read reach-level annual basin shares, multiplied them by Total_Run from AYKEscapement.xlsx, "
     "and accumulated every reach together with all of its upstream reaches through the river network."),
    ("Pairwise interannual change",
     "For all chronological year pairs, calculated signed symmetric percentage difference as "
     "200 x (later - earlier) / (later + earlier). This retains complete turn-ons (+200%) and "
     "turn-offs (-200%); zero-to-zero pairs were omitted."),
    ("Spatial scaling",
     "Related pairwise variability to accumulated upstream channel length and stream order. A "
     "distinct-tributary version retained one accumulated outlet per tributary below the basin-specific "
     "mainstem cutoff (Kuskokwim order 7; Yukon order 8)."),
    ("Absolute-production CV",
     "For each distinct tributary, calculated CV across annual fish totals and compared it with the "
     "CV of basin-wide Total_Run across the sampled years. Reported both the CV ratio and percent "
     "anomaly: 100 x (tributary CV / basin CV - 1)."),
    ("Independent-population null",
     "Simulated independent lognormal reach populations with expected production proportional to "
     "eligible channel length and assumed reach CVs of 0.25 and 1.0. Simulated production was not "
     "closed to sum to one; it was accumulated through the same network to generate reference fans "
     "and distinct-tributary CV distributions."),
]
for label, text in items:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(label + ": ")
    font(r, bold=True, color=DARK)
    font(p.add_run(text))

doc.add_heading("Results patterns worth considering", level=1)
results = [
    ("Recurring but shifting habitat modes",
     "The Kuskokwim contour panels repeatedly show several modes: low-slope/lower-distance habitat, "
     "low-slope/mid-to-upper-distance habitat, and steeper mid-distance habitat. Their relative density "
     "changes among years, suggesting redistribution among recurring habitat types more than replacement "
     "by wholly novel habitat."),
    ("Increasing spatial spread in later Kuskokwim maps",
     "The production maps appear more broadly distributed across eligible tributaries by 2021-2022 than "
     "in some earlier years. Frame this as a visual pattern unless you add a formal concentration metric."),
    ("Basin-wide CVs were similar",
     "Across sampled years, Total_Run CV was 0.244 for Kuskokwim and 0.270 for Yukon, so the basin-level "
     "difference was modest compared with the much larger differences among tributaries."),
    ("Tributary variability declined with spatial scale",
     "Median absolute-production CV generally approached the basin-wide CV as stream order increased. "
     "For Kuskokwim, the median CV ratio fell from 1.10 at order 3 to 1.02 at order 5; for Yukon, it "
     "fell from 1.36 at order 3 to 1.23 at order 7."),
    ("The Yukon signal was stronger and more heterogeneous",
     "Mean tributary CV anomaly declined from 53% above the basin-wide CV at order 3 to 21% above it "
     "at order 7. Kuskokwim anomalies were smaller (21% at order 3 and about 3-4% at orders 5-6). "
     "The Yukon interquartile ranges were also wider, indicating more heterogeneous tributary behavior."),
    ("Evidence consistent with portfolio buffering",
     "The convergence of tributary CV toward basin CV with increasing stream order is consistent with "
     "spatial aggregation buffering local variability. Phrase this as consistency with a portfolio "
     "effect, then use the independent-reach null comparison to say whether observed buffering is stronger "
     "or weaker than expected under independence."),
    ("Positive pairwise medians in the Yukon need care",
     "Many Yukon stream-order summaries have positive median signed changes, whereas Kuskokwim medians "
     "are near zero. Because the Yukon series has only four sampled years and nonconsecutive intervals, "
     "describe this as directional change among sampled years rather than a monotonic temporal trend."),
]
for label, text in results:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(label + ": ")
    font(r, bold=True, color=DARK)
    font(p.add_run(text))

doc.add_heading("Parameter checks before writing final Methods", level=1)
checks = [
    "Resolve the assignment-threshold label. params.R currently sets the main sensitivity threshold to 0.0, "
    "the contour pass is explicitly tau = 0.9, and the variance output files are labeled t0.7 even though "
    "03_VarianceBuffering.R says it reads current production outputs. Do not report a single threshold until "
    "the intended production/variance run is confirmed.",
    "Confirm whether day-stratum sampling weights should be applied to fish assignments. The scripts calculate "
    "the weights but the visible assignment loop does not use natal$weight.",
    "Describe the basin-specific error treatment accurately: the current Kuskokwim code clamps small isoscape "
    "errors to a minimum, whereas the Yukon code substitutes the basin mean isoscape error.",
    "Avoid calling the contour bands probabilities or assignment confidence intervals; they are weighted KDE "
    "quantile regions in landscape space.",
    "State the unequal time coverage and small number of annual observations explicitly (six Kuskokwim years, "
    "four Yukon years), especially for CV and temporal comparisons.",
]
for text in checks:
    p = doc.add_paragraph(style="List Bullet")
    font(p.add_run(text))

doc.add_heading("Primary files checked", level=1)
for text in [
    "Code/Analysis/params.R; 01_FullBasinRelativeProdMaps.R; 02_ContourThreshnew.R; 03_VarianceBuffering.R",
    "Outputs/PortfolioEffect/*_t0.7.csv and Figures/00_PubFigures/Fig1_KuskokwimMapsContours.png",
]:
    p = doc.add_paragraph(style="List Bullet")
    font(p.add_run(text), size=9.5, color=MUTED)

doc.save(OUT)
print(OUT)
